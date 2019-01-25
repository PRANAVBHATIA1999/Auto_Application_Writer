from io import BytesIO

import docx


def create(address, first_name, last_name, senders_class, rno, date_from, date_till, days, app_date, selected_value):
    doc = docx.Document()
    doc.add_paragraph(selected_value)
    doc.add_paragraph(address)
    doc.add_paragraph('')
    doc.add_paragraph('Date: ' + app_date)
    doc.add_paragraph('')
    doc.add_paragraph('Subject: Leave application')
    doc.add_paragraph('Respected Sir/Ma’am,')
    doc.add_paragraph(
        'I request to state that due to sudden illness I will not be able to attend school/college for ' + days + ' days '
                                                                                                                  'as the doctor has advised me to take the required amount of rest. I hope to recover soon and make up for the irregularity'
                                                                                                                  ' in studies occurred. I will return to school/college as a healthy student and take due care that my work and performance do not'
                                                                                                                  ' suffer. On the account of my sickness, I request you to kindly grant me leave from ' + date_from + ' to ' + date_till + ', I shall be utterly obliged for this.')
    doc.add_paragraph('')
    doc.add_paragraph('')
    doc.add_paragraph('')
    doc.add_paragraph('Yours obediently,')
    doc.add_paragraph(first_name + ' ' + last_name)
    doc.add_paragraph(senders_class)
    doc.add_paragraph(rno)
    f = BytesIO()
    doc.save(f)
    return f



def createVisaForm(address, first_name, last_name, senders_class, rno, date_from, date_till, app_date, purp_visit, country_name, selected_value): 
    doc = docx.Document()
    doc.add_paragraph(selected_value)
    doc.add_paragraph(address)
    doc.add_paragraph('')
    doc.add_paragraph('Date: ' + app_date)
    doc.add_paragraph('')
    doc.add_paragraph('Subject: Requesting NOC for visa application')
    doc.add_paragraph('Respected Sir/Ma’am,')
    doc.add_paragraph('I, '+ first_name + ' ' + last_name + ' roll number ' + rno + ' will be travelling to ' +
    country_name + ' with my family/friends from ' + date_from + ' till ' + date_till + '.' +  
    ' For the same, I will be applying for the required visa so I request you to kindly grant me leave and provide me with a No Objection Certificate (NOC). My purpose of visit is ' + purp_visit )
    doc.add_paragraph('')
    doc.add_paragraph('')
    doc.add_paragraph('Yours obediently,')
    doc.add_paragraph(first_name + ' ' + last_name)
    doc.add_paragraph(senders_class)
    doc.add_paragraph(rno)
    f = BytesIO()
    doc.save(f)
    return f

def createPartiForm(address, first_name, last_name, senders_class, rno, date_from, date_till, fest_name, college_name_fest, compet_name, no_of_parti, app_date, selected_value):
    doc = docx.Document()
    doc.add_paragraph(selected_value)
    doc.add_paragraph(address)
    doc.add_paragraph('')
    doc.add_paragraph('Date: ' + app_date)
    doc.add_paragraph('')
    doc.add_paragraph('Subject: Permission for participating at ' + college_name_fest)
    doc.add_paragraph('Respected Sir/Ma’am,')
    doc.add_paragraph('I, ' + first_name + ' ' + last_name + ' the student ambassador, would like to bring to your notice that we propose to go to ' 
    + compet_name + ' the annual fest of ' + college_name_fest + '. We will be a team of ' + no_of_parti + ' members and will be going from ' +
    date_from + ' to ' + date_till + 'we will be participating in ' + fest_name + ' I request you to kindly grant us permission and attendance for the same.')
    doc.add_paragraph('')
    doc.add_paragraph('')
    doc.add_paragraph('Yours obediently,')
    doc.add_paragraph(first_name + ' ' + last_name)
    doc.add_paragraph(senders_class)
    doc.add_paragraph(rno)
    f = BytesIO()
    doc.save(f)
    return f

